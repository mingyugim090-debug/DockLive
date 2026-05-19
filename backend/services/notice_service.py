import html
import json
import re
import shutil
import sys
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any

from core.errors import AnalysisError
from models.schemas import NoticeDocument, NoticeSection
from services.ai_provider import call_json, should_use_mock_ai
from services.document_ingestion import convert_hwp_to_hwpx, extract_text_from_hwpx
from services.drafting_service import _make_tmp_dir, _require_hwpx_scripts, _run_hwpx_command, _safe_title
from services.pdf_export_service import convert_hwpx_bytes_to_pdf


FORBIDDEN_NOTICE_TERMS = [
    "LiveDock: 자동작성 내용",
    "DockLive: 자동작성 내용",
    "자동작성 초안",
    "요약",
    "JSON 문자열",
    "추가 입력 필요",
    "내부 상태값",
    "generated_fields",
    "documentType",
    "template_id",
    "validation_summary",
    "sections",
    "contact",
]

CANONICAL_NOTICE_SECTION_HEADINGS = [
    "사업 개요",
    "모집 대상",
    "모집 인원",
    "신청 기간",
    "운영 일정",
    "신청 방법",
    "선정 기준",
    "제출 서류",
]

NOTICE_TEMPLATES: dict[str, dict[str, Any]] = {
    "startup_camp_notice": {
        "name": "창업캠프 모집 공고문",
        "purpose": "창업캠프 참가자 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "business_support_notice": {
        "name": "지원사업 참여기업 모집 공고문",
        "purpose": "지원사업 참여기업 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "education_program_notice": {
        "name": "교육 프로그램 수강생 모집 공고문",
        "purpose": "교육 프로그램 수강생 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "event_participant_notice": {
        "name": "행사 참가자 모집 공고문",
        "purpose": "행사 참가자 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "scholarship_notice": {
        "name": "장학생 모집 공고문",
        "purpose": "장학생 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "tenant_company_notice": {
        "name": "입주기업 모집 공고문",
        "purpose": "창업보육센터·공간 입주기업 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "research_participant_notice": {
        "name": "연구과제 참여자 모집 공고문",
        "purpose": "연구과제 참여자 모집",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
    "bid_rfp_notice": {
        "name": "입찰/제안요청 공고문",
        "purpose": "입찰 또는 제안서 접수",
        "sections": CANONICAL_NOTICE_SECTION_HEADINGS,
    },
}


def generate_notice_document(
    template_id: str,
    inputs: dict[str, str],
    reference_documents: list[dict[str, str]] | None = None,
) -> tuple[NoticeDocument, list[str]]:
    template = NOTICE_TEMPLATES.get(template_id)
    if template is None:
        raise AnalysisError(f"알 수 없는 공고문 템플릿입니다: {template_id}")

    cleaned_inputs = {str(key): str(value).strip() for key, value in inputs.items() if str(value).strip()}
    references = reference_documents or []
    warnings: list[str] = []

    if should_use_mock_ai():
        data = _mock_notice_document(template_id, template, cleaned_inputs, references)
    else:
        data = _call_notice_ai(template_id, template, cleaned_inputs, references)

    document = _normalize_notice_document(template_id, template, data, cleaned_inputs)
    warnings.extend(_quality_warnings(document))
    _assert_clean_notice_text(render_notice_markdown(document), "공고문 미리보기")
    return document, warnings


def render_notice_markdown(document: NoticeDocument) -> str:
    lines = [
        f"# {document.title}",
        "",
        document.organization,
        "",
        "공고 안내문",
        f"{document.organization}은(는) {document.purpose}을(를) 다음과 같이 공고하오니 관심 있는 대상자의 많은 신청 바랍니다.",
        "",
    ]

    for index, section in enumerate(document.sections, start=1):
        heading = section.heading.strip()
        if not re.match(r"^\d+\.", heading):
            heading = f"{index}. {heading}"
        lines.extend([f"## {heading}", section.body.strip(), ""])

    lines.extend(
        [
            "## 9. 문의처",
            f"- 담당 부서: {_cell(document.contact.department)}",
            f"- 연락처: {_cell(document.contact.phone)}",
            f"- 이메일: {_cell(document.contact.email)}",
            "",
            "## 붙임 문서 목록",
        ]
    )
    if document.attachments:
        lines.extend(f"{idx}. {item}" for idx, item in enumerate(document.attachments, start=1))
    else:
        lines.append("해당 없음")
    return "\n".join(lines).strip() + "\n"


def export_notice_hwpx(document: NoticeDocument) -> tuple[str, bytes, dict[str, Any]]:
    markdown = render_notice_markdown(document)
    _assert_clean_notice_text(markdown, "HWPX 입력")
    filename, content, summary = _export_notice_document_to_hwpx(document)
    excerpt = str(summary.get("extracted_text_excerpt") or "")
    _assert_clean_notice_text(excerpt or markdown, "HWPX 추출 텍스트")
    summary["renderer"] = "notice_document_renderer"
    summary["forbidden_terms_checked"] = FORBIDDEN_NOTICE_TERMS
    return filename, content, summary


def _export_notice_document_to_hwpx(document: NoticeDocument) -> tuple[str, bytes, dict[str, Any]]:
    scripts = _require_hwpx_scripts(
        "build_hwpx.py",
        "fix_namespaces.py",
        "validate.py",
        "text_extract.py",
        "verify_hwpx.py",
        "hwpx_helpers.py",
    )
    tmpdir = _make_tmp_dir()
    safe_title = _safe_title(document.title)
    summary: dict[str, Any] = {
        "validation_passed": False,
        "namespace_fixed": False,
        "warnings": [],
        "generation_method": "notice-document-xml-renderer",
    }
    try:
        section_path = tmpdir / "section0.xml"
        output_path = tmpdir / f"{safe_title}.hwpx"
        verify_path = tmpdir / "verify.json"
        section_path.write_text(_build_notice_section_xml(document), encoding="utf-8")

        python_bin = sys.executable or "python"
        _run_hwpx_command(
            [
                python_bin,
                str(scripts["build_hwpx.py"]),
                "--section",
                str(section_path),
                "--title",
                document.title,
                "--creator",
                document.organization or "DockLive",
                "--output",
                str(output_path),
            ]
        )
        fix_result = _run_hwpx_command([python_bin, str(scripts["fix_namespaces.py"]), str(output_path)])
        summary["namespace_fixed"] = True
        summary["namespace_output"] = (fix_result.stdout or fix_result.stderr or "").strip()[:1000]
        validate_result = _run_hwpx_command([python_bin, str(scripts["validate.py"]), str(output_path)])
        summary["validation_passed"] = True
        summary["validation_output"] = (validate_result.stdout or validate_result.stderr or "").strip()[:1000]

        try:
            _run_hwpx_command(
                [
                    python_bin,
                    str(scripts["verify_hwpx.py"]),
                    "--result",
                    str(output_path),
                    "--json",
                    str(verify_path),
                ]
            )
            summary["verify_passed"] = True
            if verify_path.exists():
                summary["verify_report"] = json.loads(verify_path.read_text(encoding="utf-8"))
        except Exception as exc:
            summary["verify_passed"] = False
            summary["warnings"].append(f"verify_hwpx.py 확인을 완료하지 못했습니다: {str(exc)[:300]}")

        try:
            text_result = _run_hwpx_command([python_bin, str(scripts["text_extract.py"]), str(output_path), "--include-tables"])
            extracted = text_result.stdout or ""
            summary["text_extract_passed"] = True
            summary["text_chars"] = len(extracted)
            summary["title_found"] = bool(document.title and document.title[:20] in extracted)
            summary["extracted_text_excerpt"] = extracted[:1000]
        except Exception as exc:
            summary["text_extract_passed"] = False
            summary["text_chars"] = 0
            summary["warnings"].append(f"text_extract.py 확인을 완료하지 못했습니다: {str(exc)[:300]}")

        if not output_path.exists() or output_path.stat().st_size < 1000:
            raise AnalysisError("공고문 HWPX 파일이 정상적으로 생성되지 않았습니다.")
        return output_path.name, output_path.read_bytes(), summary
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _build_notice_section_xml(document: NoticeDocument) -> str:
    scripts_dir = Path(__file__).resolve().parents[1] / "hwpx_toolchain" / "scripts"
    if str(scripts_dir) not in sys.path:
        sys.path.insert(0, str(scripts_dir))
    from hwpx_helpers import (
        NS_DECL,
        extract_secpr_and_colpr,
        make_empty_line,
        make_first_para,
        make_text_para,
        next_id,
        reset_id,
        xml_escape,
    )

    reference = Path(__file__).resolve().parents[1] / "hwpx_toolchain" / "templates" / "reference.hwpx"
    secpr, colpr = extract_secpr_and_colpr(reference)
    reset_id()

    def make_table(
        headers: list[str],
        rows: list[list[str]],
        col_widths: list[int] | None = None,
        row_height: int = 2300,
    ) -> str:
        num_cols = max(1, len(headers))
        body_width = 42520
        if col_widths is None:
            base_width = body_width // num_cols
            col_widths = [base_width] * num_cols
            col_widths[-1] += body_width - sum(col_widths)
        total_rows = 1 + len(rows)
        p_id = next_id()
        tbl_id = next_id()

        def make_cell(text: str, is_header: bool, col_idx: int, row_idx: int) -> str:
            cell_pid = next_id()
            border_fill = "4" if is_header else "3"
            char_pr = "25" if is_header else "20"
            para_pr = "25" if is_header else "24"
            return (
                f'<hp:tc name="" header="{1 if is_header else 0}" hasMargin="0" protect="0" editable="0" dirty="1" borderFillIDRef="{border_fill}">'
                f'<hp:subList id="" textDirection="HORIZONTAL" lineWrap="BREAK" vertAlign="CENTER" '
                f'linkListIDRef="0" linkListNextIDRef="0" textWidth="0" textHeight="0" hasTextRef="0" hasNumRef="0">'
                f'<hp:p paraPrIDRef="{para_pr}" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0" id="{cell_pid}">'
                f'<hp:run charPrIDRef="{char_pr}"><hp:t>{xml_escape(text)}</hp:t></hp:run>'
                f'</hp:p></hp:subList>'
                f'<hp:cellAddr colAddr="{col_idx}" rowAddr="{row_idx}"/>'
                f'<hp:cellSpan colSpan="1" rowSpan="1"/>'
                f'<hp:cellSz width="{col_widths[col_idx]}" height="{row_height}"/>'
                f'<hp:cellMargin left="220" right="220" top="120" bottom="120"/>'
                f'</hp:tc>'
            )

        table_rows = [
            "<hp:tr>" + "".join(make_cell(header, True, col_idx, 0) for col_idx, header in enumerate(headers)) + "</hp:tr>"
        ]
        for row_idx, row in enumerate(rows, start=1):
            padded = row[:num_cols] + [""] * max(0, num_cols - len(row))
            table_rows.append(
                "<hp:tr>"
                + "".join(make_cell(padded[col_idx], False, col_idx, row_idx) for col_idx in range(num_cols))
                + "</hp:tr>"
            )

        return (
            f'<hp:p id="{p_id}" paraPrIDRef="0" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0">'
            f'<hp:run charPrIDRef="0">'
            f'<hp:tbl id="{tbl_id}" zOrder="0" numberingType="TABLE" textWrap="TOP_AND_BOTTOM" '
            f'textFlow="BOTH_SIDES" lock="0" dropcapstyle="None" pageBreak="CELL" repeatHeader="0" '
            f'rowCnt="{total_rows}" colCnt="{num_cols}" cellSpacing="0" borderFillIDRef="3" noAdjust="0">'
            f'<hp:sz width="{sum(col_widths)}" widthRelTo="ABSOLUTE" height="{row_height * total_rows}" heightRelTo="ABSOLUTE" protect="0"/>'
            f'<hp:pos treatAsChar="1" affectLSpacing="0" flowWithText="1" allowOverlap="0" '
            f'holdAnchorAndSO="0" vertRelTo="PARA" horzRelTo="COLUMN" vertAlign="TOP" horzAlign="LEFT" vertOffset="0" horzOffset="0"/>'
            f'<hp:outMargin left="0" right="0" top="0" bottom="0"/>'
            f'<hp:inMargin left="0" right="0" top="0" bottom="0"/>'
            f'{"".join(table_rows)}'
            f'</hp:tbl></hp:run></hp:p>'
        )

    def first_body(keyword: str, fallback: str = "") -> str:
        for section in document.sections:
            if keyword in section.heading:
                return " ".join(_split_body_lines(section.body)) or fallback
        return fallback

    overview_rows = [
        ["공고 유형", document.purpose or document.documentType],
        ["주관 기관", document.organization],
        ["신청 기간", document.schedule.applicationPeriod or "공고문 참조"],
        ["운영 기간", document.schedule.eventPeriod or "선정 이후 별도 안내"],
        ["접수 방법", document.applicationMethod or "붙임 서식 작성 후 제출"],
    ]
    application_rows = [
        ["신청자/기업명", "OOO"],
        ["소속/대표자", document.organization],
        ["신청 분야", document.purpose or document.documentType],
        ["연락처", f"{_cell(document.contact.phone)} / {_cell(document.contact.email)}"],
    ]
    schedule_rows = [
        ["1", "공고 및 접수", document.schedule.applicationPeriod or "공고문 참조", "신청서 및 붙임 서류 접수"],
        ["2", "요건 검토", "접수 마감 후", "자격 요건 및 제출 서류 확인"],
        ["3", "평가 및 선정", "별도 안내", "평가 기준에 따른 심사"],
        ["4", "결과 안내", "선정 후", "홈페이지 또는 개별 통보"],
    ]
    evaluation_rows = [
        ["적합성", "30%", first_body("목적", "공고 목적과 신청 내용의 부합 정도")],
        ["실현 가능성", "30%", first_body("지원", "추진 계획, 일정, 수행 역량의 구체성")],
        ["기대 효과", "25%", first_body("평가", "성과 확산 가능성 및 공공성")],
        ["서류 완성도", "15%", first_body("서류", "제출 서류의 충실도와 사실 확인 가능성")],
    ]

    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>',
        f"<hs:sec {NS_DECL}>",
        make_first_para(secpr, colpr, charpr="8", parapr="24"),
        make_text_para(document.title, charpr="7", parapr="19"),
        make_empty_line(charpr="8", parapr="24"),
        make_text_para(document.organization, charpr="19", parapr="20"),
        make_empty_line(charpr="8", parapr="24"),
        make_text_para("공고 안내문", charpr="25", parapr="25"),
        make_text_para(
            f"{document.organization}은(는) {document.purpose}을(를) 다음과 같이 공고하오니 관심 있는 대상자의 많은 신청 바랍니다.",
            charpr="20",
            parapr="24",
        ),
        make_empty_line(charpr="8", parapr="24"),
        make_table(["구분", "내용"], overview_rows, col_widths=[10500, 32020]),
        make_empty_line(charpr="8", parapr="24"),
        make_text_para("참가 신청서", charpr="25", parapr="25"),
        make_table(["항목", "작성 내용"], application_rows, col_widths=[12000, 30520]),
        make_text_para("신청 내용 및 추진 계획", charpr="25", parapr="25"),
        make_table(
            ["구분", "주요 내용"],
            [
                ["신청 내용", first_body("개요", document.purpose or "공고 목적에 맞는 신청 내용을 작성합니다.")],
                ["추진 계획", first_body("방법", document.applicationMethod or "접수 후 선정 절차에 따라 추진합니다.")],
            ],
            col_widths=[12000, 30520],
            row_height=3000,
        ),
        make_empty_line(charpr="8", parapr="24"),
        make_text_para("추진 일정", charpr="25", parapr="25"),
        make_table(["단계", "절차", "일정", "확인 사항"], schedule_rows, col_widths=[5000, 9500, 12000, 16020]),
        make_empty_line(charpr="8", parapr="24"),
        make_text_para("평가 기준", charpr="25", parapr="25"),
        make_table(["평가 항목", "비중", "세부 기준"], evaluation_rows, col_widths=[8500, 6500, 27520]),
    ]

    for section in document.sections:
        parts.append(make_empty_line(charpr="8", parapr="24"))
        parts.append(make_text_para(section.heading, charpr="25", parapr="25"))
        for line in _split_body_lines(section.body):
            parts.append(make_text_para(line, charpr="20", parapr="24"))

    parts.append(make_empty_line(charpr="8", parapr="24"))
    parts.append(make_text_para("9. 문의처", charpr="25", parapr="25"))
    parts.append(
        make_table(
            ["담당 부서", "연락처", "이메일"],
            [[_cell(document.contact.department), _cell(document.contact.phone), _cell(document.contact.email)]],
            col_widths=[13000, 12500, 17020],
        )
    )

    parts.append(make_empty_line(charpr="8", parapr="24"))
    parts.append(make_text_para("붙임 문서 목록", charpr="25", parapr="25"))
    attachments = document.attachments or ["해당 없음"]
    parts.append(
        make_table(
            ["번호", "문서명", "비고"],
            [[str(index), item, "1부"] for index, item in enumerate(attachments, start=1)],
            col_widths=[5000, 28520, 9000],
        )
    )

    parts.append("</hs:sec>")
    return "\n".join(parts)


def _export_notice_markdown_to_hwpx(markdown: str, title: str) -> tuple[str, bytes, dict[str, Any]]:
    scripts = _require_hwpx_scripts("md2hwpx.py", "fix_namespaces.py", "validate.py", "text_extract.py", "verify_hwpx.py")
    tmpdir = _make_tmp_dir()
    safe_title = _safe_title(title)
    summary: dict[str, Any] = {
        "validation_passed": False,
        "namespace_fixed": False,
        "warnings": [],
        "generation_method": "notice-md2hwpx.py",
    }
    try:
        markdown_path = tmpdir / "notice.md"
        output_path = tmpdir / f"{safe_title}.hwpx"
        verify_path = tmpdir / "verify.json"
        markdown_path.write_text(markdown, encoding="utf-8")

        python_bin = sys.executable or "python"
        _run_hwpx_command([python_bin, str(scripts["md2hwpx.py"]), str(markdown_path), "-o", str(output_path)])
        fix_result = _run_hwpx_command([python_bin, str(scripts["fix_namespaces.py"]), str(output_path)])
        summary["namespace_fixed"] = True
        summary["namespace_output"] = (fix_result.stdout or fix_result.stderr or "").strip()[:1000]
        validate_result = _run_hwpx_command([python_bin, str(scripts["validate.py"]), str(output_path)])
        summary["validation_passed"] = True
        summary["validation_output"] = (validate_result.stdout or validate_result.stderr or "").strip()[:1000]

        try:
            _run_hwpx_command(
                [
                    python_bin,
                    str(scripts["verify_hwpx.py"]),
                    "--result",
                    str(output_path),
                    "--json",
                    str(verify_path),
                ]
            )
            summary["verify_passed"] = True
            if verify_path.exists():
                summary["verify_report"] = json.loads(verify_path.read_text(encoding="utf-8"))
        except Exception as exc:
            summary["verify_passed"] = False
            summary["warnings"].append(f"verify_hwpx.py 확인을 완료하지 못했습니다: {str(exc)[:300]}")

        try:
            text_result = _run_hwpx_command([python_bin, str(scripts["text_extract.py"]), str(output_path), "--include-tables"])
            extracted = text_result.stdout or ""
            summary["text_extract_passed"] = True
            summary["text_chars"] = len(extracted)
            summary["title_found"] = bool(title and title[:20] in extracted)
            summary["extracted_text_excerpt"] = extracted[:1000]
        except Exception as exc:
            summary["text_extract_passed"] = False
            summary["text_chars"] = 0
            summary["warnings"].append(f"text_extract.py 확인을 완료하지 못했습니다: {str(exc)[:300]}")

        if not output_path.exists() or output_path.stat().st_size < 1000:
            raise AnalysisError("공고문 HWPX 파일이 정상적으로 생성되지 않았습니다.")
        return output_path.name, output_path.read_bytes(), summary
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def export_notice_pdf(document: NoticeDocument) -> tuple[str, bytes, dict[str, Any]]:
    hwpx_filename, hwpx_content, hwpx_summary = export_notice_hwpx(document)
    pdf_filename, pdf_content, pdf_summary = convert_hwpx_bytes_to_pdf(
        hwpx_content,
        document.title,
        source_filename=hwpx_filename,
    )
    return pdf_filename, pdf_content, {"hwpx_validation": hwpx_summary, "pdf_validation": pdf_summary}


def export_notice_docx(document: NoticeDocument) -> tuple[str, bytes, dict[str, Any]]:
    try:
        content = _build_docx_with_python_docx(document)
        method = "python-docx"
    except Exception as exc:
        content = _build_minimal_docx(document)
        method = f"minimal-openxml-fallback: {exc.__class__.__name__}"
    safe_title = _safe_filename(document.title)
    return f"{safe_title}.docx", content, {"renderer": "notice_document_renderer", "method": method}


def extract_reference_text(content: bytes, filename: str) -> tuple[str, list[str]]:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        from services.pdf_parser import extract_text_from_pdf

        return extract_text_from_pdf(content, filename), []
    if suffix == ".hwpx":
        return extract_text_from_hwpx(content)
    if suffix == ".hwp":
        converted, warnings = convert_hwp_to_hwpx(content, filename)
        text, extract_warnings = extract_text_from_hwpx(converted)
        return text, warnings + extract_warnings
    if suffix in {".txt", ".md"}:
        return content.decode("utf-8", errors="replace"), []
    if suffix == ".docx":
        return _extract_text_from_docx(content), []
    raise AnalysisError("참고자료는 PDF, DOCX, HWP, HWPX, TXT, MD 파일만 업로드할 수 있습니다.")


def _call_notice_ai(
    template_id: str,
    template: dict[str, Any],
    inputs: dict[str, str],
    references: list[dict[str, str]],
) -> dict[str, Any]:
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "documentType": {"type": "string"},
            "title": {"type": "string"},
            "organization": {"type": "string"},
            "purpose": {"type": "string"},
            "applicationMethod": {"type": "string"},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "heading": {"type": "string"},
                        "body": {"type": "string"},
                    },
                    "required": ["heading", "body"],
                },
            },
            "schedule": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "applicationPeriod": {"type": "string"},
                    "eventPeriod": {"type": "string"},
                },
                "required": ["applicationPeriod", "eventPeriod"],
            },
            "contact": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "department": {"type": "string"},
                    "phone": {"type": "string"},
                    "email": {"type": "string"},
                },
                "required": ["department", "phone", "email"],
            },
            "attachments": {"type": "array", "items": {"type": "string"}},
        },
        "required": [
            "documentType",
            "title",
            "organization",
            "purpose",
            "applicationMethod",
            "sections",
            "schedule",
            "contact",
            "attachments",
        ],
    }
    system_prompt = (
        "당신은 한국 공공기관 공고문을 작성하는 행정 문서 전문가입니다. "
        "내부 메타데이터, JSON 문자열, 자동작성 문구를 본문에 쓰지 말고 제출 가능한 공고문 내용만 작성하세요. "
        "불확실한 사실은 임의로 확정하지 말고 확인 필요 문장으로 완곡하게 작성하세요."
    )
    reference_excerpt = "\n\n".join(
        f"[{item.get('filename', '참고자료')}]\n{item.get('text', '')[:2500]}" for item in references
    )[:8000]
    payload = {
        "templateId": template_id,
        "templateName": template["name"],
        "purpose": template["purpose"],
        "requiredSections": template["sections"],
        "inputs": inputs,
        "referenceExcerpt": reference_excerpt,
    }
    user_prompt = "다음 정보를 공공기관 공고문 document JSON으로 작성하세요.\n" + json.dumps(payload, ensure_ascii=False)
    return call_json("draft", system_prompt, user_prompt, max_tokens=4096, json_schema=schema, schema_name="notice_document")


def _normalize_notice_document(
    template_id: str,
    template: dict[str, Any],
    raw: dict[str, Any],
    inputs: dict[str, str],
) -> NoticeDocument:
    document = NoticeDocument.model_validate(
        {
            "documentType": template_id,
            "title": raw.get("title") or inputs.get("title") or template["name"],
            "organization": raw.get("organization") or inputs.get("organization") or "기관명 확인 필요",
            "purpose": raw.get("purpose") or template["purpose"],
            "applicationMethod": raw.get("applicationMethod") or inputs.get("applicationMethod") or "온라인 또는 담당 부서 접수",
            "sections": raw.get("sections") or [],
            "schedule": raw.get("schedule") or {},
            "contact": raw.get("contact") or {},
            "attachments": raw.get("attachments") or [],
        }
    )
    if not document.sections:
        document.sections = [
            NoticeSection(heading=f"{index}. {heading}", body=_default_section_body(heading, inputs))
            for index, heading in enumerate(template["sections"], start=1)
        ]
    body_by_heading: dict[str, str] = {}
    for section in document.sections:
        normalized_heading = _canonical_heading(section.heading)
        body = section.body.strip()
        if normalized_heading and body:
            body_by_heading[normalized_heading] = body

    normalized_sections = []
    for index, heading in enumerate(CANONICAL_NOTICE_SECTION_HEADINGS, start=1):
        body = body_by_heading.get(heading) or _default_section_body(heading, inputs)
        normalized_sections.append(NoticeSection(heading=f"{index}. {heading}", body=body))
    document.sections = normalized_sections
    if not document.attachments:
        document.attachments = _split_attachments(inputs.get("attachments")) or _split_attachments(inputs.get("documents")) or ["신청서", "개인정보 수집 및 이용 동의서"]
    return document


def _mock_notice_document(
    template_id: str,
    template: dict[str, Any],
    inputs: dict[str, str],
    references: list[dict[str, str]],
) -> dict[str, Any]:
    title = inputs.get("title") or template["name"]
    organization = inputs.get("organization") or "서울과학기술대학교 창업지원단"
    target = inputs.get("target") or "해당 분야에 관심 있는 시민 및 기관 관계자"
    capacity = inputs.get("capacity") or "모집 규모는 기관 계획과 예산 범위에 따라 정합니다."
    benefit = inputs.get("benefit") or "전문 교육, 네트워킹, 후속 지원 기회"
    documents = inputs.get("documents") or inputs.get("attachments") or "신청서, 개인정보 수집 및 이용 동의서 등 공고에서 정한 서류"
    criteria = inputs.get("selectionCriteria") or "신청 자격, 사업 목적 적합성, 제출 서류의 충실성 등을 종합적으로 검토합니다."
    reference_note = " 참고자료의 주요 내용을 반영하여 세부 일정과 운영 방식은 담당 부서 확인 후 확정합니다." if references else ""
    return {
        "documentType": template_id,
        "title": title,
        "organization": organization,
        "purpose": template["purpose"],
        "applicationMethod": inputs.get("applicationMethod") or "기관 누리집 공고문 확인 후 온라인 신청",
        "sections": [
            {
                "heading": f"{idx}. {heading}",
                "body": _default_section_body(
                    heading,
                    inputs,
                    target=target,
                    capacity=capacity,
                    benefit=benefit,
                    documents=documents,
                    criteria=criteria,
                ) + reference_note,
            }
            for idx, heading in enumerate(template["sections"], start=1)
        ],
        "schedule": {
            "applicationPeriod": inputs.get("applicationPeriod") or "공고일로부터 2026. 6. 30.까지",
            "eventPeriod": inputs.get("eventPeriod") or inputs.get("operationPeriod") or "선정 후 별도 안내",
        },
        "contact": {
            "department": inputs.get("department") or "사업 담당 부서",
            "phone": inputs.get("phone") or "02-000-0000",
            "email": inputs.get("email") or "notice@example.go.kr",
        },
        "attachments": _split_attachments(inputs.get("attachments")) or ["신청서", "개인정보 수집 및 이용 동의서"],
    }


def _default_section_body(
    heading: str,
    inputs: dict[str, str],
    target: str = "공고 대상자",
    capacity: str = "모집 규모는 기관 계획과 예산 범위에 따라 정합니다.",
    benefit: str = "세부 지원 내용",
    documents: str = "신청서, 개인정보 수집 및 이용 동의서 등 공고에서 정한 서류",
    criteria: str = "신청 자격, 사업 목적 적합성, 제출 서류의 충실성 등을 종합적으로 검토합니다.",
) -> str:
    title = inputs.get("title") or "본 공고"
    organization = inputs.get("organization") or "주관 기관"
    application_period = inputs.get("applicationPeriod") or "공고문에 따른 접수 기간"
    event_period = inputs.get("eventPeriod") or "선정 이후 별도 안내하는 운영 일정"
    method = inputs.get("applicationMethod") or "기관 누리집 또는 담당 부서가 안내한 접수 방법"
    if heading == "사업 개요":
        return f"{organization}은 {title}을 통해 사업 목적에 적합한 참여자를 모집하고, 원활한 운영을 위한 절차를 안내합니다. 주요 지원 내용은 {benefit}입니다."
    if heading == "모집 대상":
        return f"모집 대상은 {target}이며, 세부 자격 요건은 공고문과 제출 서류를 기준으로 확인합니다."
    if heading == "모집 인원":
        return capacity
    if heading == "신청 기간":
        return application_period
    if heading == "운영 일정":
        return f"운영 일정은 {event_period}이며, 신청 접수, 심사, 선정 안내, 사업 운영 순으로 진행합니다."
    if heading == "신청 방법":
        return f"신청자는 제출 서류를 준비하여 {method}으로 접수해야 합니다."
    if heading == "선정 기준":
        return criteria
    if heading == "제출 서류":
        return documents
    return "세부 사항은 공고문 본문과 붙임 서류를 확인하여 주시기 바랍니다."


def _quality_warnings(document: NoticeDocument) -> list[str]:
    warnings = []
    if not document.contact.phone or not document.contact.email:
        warnings.append("문의처 일부가 비어 있습니다. 다운로드 전 실제 연락처를 확인해 주세요.")
    return warnings


def _assert_clean_notice_text(text: str, label: str) -> None:
    for term in FORBIDDEN_NOTICE_TERMS:
        if term in text:
            raise AnalysisError(f"{label}에 내부 문구가 포함되어 export를 중단했습니다: {term}")
    if re.search(r"\{\s*\"(?:documentType|sections|generated_fields)\"", text):
        raise AnalysisError(f"{label}에 JSON 문자열이 포함되어 export를 중단했습니다.")


def _cell(value: str) -> str:
    return (value or "-").replace("|", "/").replace("\n", "<br>")


def _split_body_lines(value: str) -> list[str]:
    clean = re.sub(r"\r\n?", "\n", value or "").strip()
    if not clean:
        return ["-"]
    lines = [line.strip(" \t-") for line in clean.split("\n") if line.strip(" \t-")]
    return lines or [clean]


def _canonical_heading(value: str) -> str:
    heading = re.sub(r"^\d+\.\s*", "", (value or "").strip())
    heading = heading.replace("지원 대상", "모집 대상").replace("신청 자격", "모집 대상")
    for canonical in CANONICAL_NOTICE_SECTION_HEADINGS:
        if canonical == heading or canonical in heading:
            return canonical
    if "개요" in heading:
        return "사업 개요"
    if "대상" in heading or "자격" in heading:
        return "모집 대상"
    if "인원" in heading or "규모" in heading:
        return "모집 인원"
    if "신청 기간" in heading or "접수 기간" in heading:
        return "신청 기간"
    if "일정" in heading or "기간" in heading:
        return "운영 일정"
    if "방법" in heading or "접수" in heading:
        return "신청 방법"
    if "선정" in heading or "평가" in heading:
        return "선정 기준"
    if "서류" in heading:
        return "제출 서류"
    return ""


def _split_attachments(value: str | None) -> list[str]:
    if not value:
        return []
    return [item.strip(" -\t") for item in re.split(r"[\n,]+", value) if item.strip(" -\t")]


def _safe_filename(title: str) -> str:
    safe = re.sub(r'[\\/:*?"<>|]+', "_", title).strip().strip(".")
    return safe[:80] or "notice_document"


def _extract_text_from_docx(content: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(content), "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    except Exception as exc:
        raise AnalysisError("DOCX 참고자료에서 텍스트를 추출하지 못했습니다.") from exc
    text = re.sub(r"<w:tab[^>]*/>", "\t", xml)
    text = re.sub(r"</w:p>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return html.unescape(text).strip()


def _build_docx_with_python_docx(document: NoticeDocument) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading(document.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "구분"
    table.rows[0].cells[1].text = "내용"
    for label, value in [
        ("기관명", document.organization),
        ("공고 유형", document.purpose),
        ("신청 기간", document.schedule.applicationPeriod),
        ("운영 기간", document.schedule.eventPeriod),
        ("접수 방법", document.applicationMethod),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value or "확인 필요"
    for section in document.sections:
        doc.add_heading(section.heading, level=1)
        doc.add_paragraph(section.body)
    doc.add_heading("문의처", level=1)
    contact = doc.add_table(rows=2, cols=3)
    contact.style = "Table Grid"
    contact.rows[0].cells[0].text = "부서"
    contact.rows[0].cells[1].text = "연락처"
    contact.rows[0].cells[2].text = "이메일"
    contact.rows[1].cells[0].text = document.contact.department or "-"
    contact.rows[1].cells[1].text = document.contact.phone or "-"
    contact.rows[1].cells[2].text = document.contact.email or "-"
    doc.add_heading("붙임", level=1)
    for item in document.attachments or ["해당 없음"]:
        doc.add_paragraph(item, style="List Number")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _build_minimal_docx(document: NoticeDocument) -> bytes:
    paragraphs = [document.title, document.organization, document.purpose]
    paragraphs.extend(f"{section.heading}\n{section.body}" for section in document.sections)
    paragraphs.append(f"문의처: {document.contact.department} / {document.contact.phone} / {document.contact.email}")
    paragraphs.append("붙임: " + ", ".join(document.attachments or ["해당 없음"]))
    body = "".join(f"<w:p><w:r><w:t>{html.escape(text)}</w:t></w:r></w:p>" for text in paragraphs)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}<w:sectPr/></w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )
    output = BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
    return output.getvalue()
