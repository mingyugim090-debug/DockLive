"""Workspace document exports (Markdown/HTML/DOCX/HWPX/PDF).

Charts fall back to their source table in every export format — native HWPX
chart objects stay out of scope per the table-first contract.
"""

import html as html_lib
import io
import re
from typing import Any

from models.schemas import ChartSpec, GeneratedDocument, ParsedTableCell, VisualBlock


def _chart_fallback_rows(chart: ChartSpec) -> list[list[str]]:
    header = ["구분", *[series.name for series in chart.series]]
    rows = [header]
    for index, label in enumerate(chart.labels):
        row = [label]
        for series in chart.series:
            value = series.values[index] if index < len(series.values) else ""
            row.append(f"{value:g}" if isinstance(value, float) else str(value))
        rows.append(row)
    return rows


def _rows_as_text(rows: list[list[ParsedTableCell]]) -> list[list[str]]:
    return [[cell.text for cell in row] for row in rows]


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(normalized[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


def _block_markdown(block: VisualBlock) -> str:
    if block.kind == "heading":
        return f"## {block.markdown}".strip()
    if block.kind == "paragraph":
        return block.markdown.strip()
    if block.kind == "table":
        return _markdown_table(_rows_as_text(block.rows))
    if block.kind == "chart" and block.chart:
        note = f"> 그래프: {block.chart.title or '차트'} — 웹 미리보기 전용, 내보내기에서는 표로 표시됩니다."
        return note + "\n\n" + _markdown_table(_chart_fallback_rows(block.chart))
    return ""


def render_markdown(document: GeneratedDocument) -> str:
    parts = [f"# {document.title}".strip()]
    parts.extend(part for part in (_block_markdown(block) for block in document.blocks) if part)
    return "\n\n".join(parts).strip() + "\n"


def _html_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    head_cells = "".join(f"<th>{html_lib.escape(cell)}</th>" for cell in rows[0])
    body_rows = "".join(
        "<tr>" + "".join(f"<td>{html_lib.escape(cell)}</td>" for cell in row) + "</tr>"
        for row in rows[1:]
    )
    return f"<table><thead><tr>{head_cells}</tr></thead><tbody>{body_rows}</tbody></table>"


def _block_html(block: VisualBlock) -> str:
    if block.kind == "heading":
        return f"<h2>{html_lib.escape(block.markdown)}</h2>"
    if block.kind == "paragraph":
        lines = [html_lib.escape(line) for line in block.markdown.splitlines() if line.strip()]
        return "<p>" + "<br/>".join(lines) + "</p>" if lines else ""
    if block.kind == "table":
        return _html_table(_rows_as_text(block.rows))
    if block.kind == "chart" and block.chart:
        caption = html_lib.escape(block.chart.title or "차트")
        return (
            f"<p><em>그래프: {caption} — 웹 미리보기 전용, 내보내기에서는 표로 표시됩니다.</em></p>"
            + _html_table(_chart_fallback_rows(block.chart))
        )
    return ""


def render_html(document: GeneratedDocument) -> str:
    body = "\n".join(part for part in (_block_html(block) for block in document.blocks) if part)
    style = (
        "body{font-family:'Malgun Gothic',sans-serif;max-width:760px;margin:40px auto;color:#24312D;}"
        "table{border-collapse:collapse;width:100%;margin:12px 0;}"
        "th,td{border:1px solid #DDE7E2;padding:6px 10px;font-size:14px;text-align:left;}"
        "th{background:#EDF7F2;color:#245D50;}h1,h2{color:#245D50;}"
    )
    return (
        "<!DOCTYPE html><html lang=\"ko\"><head><meta charset=\"utf-8\">"
        f"<title>{html_lib.escape(document.title)}</title><style>{style}</style></head>"
        f"<body><h1>{html_lib.escape(document.title)}</h1>\n{body}\n</body></html>"
    )


def _safe_filename(title: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n]+', "_", (title or "document").strip())
    return cleaned[:80] or "document"


def _docx_add_table(docx_document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = docx_document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell_text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.text = cell_text
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def export_docx(document: GeneratedDocument) -> tuple[str, bytes]:
    from docx import Document as DocxDocument

    docx_document = DocxDocument()
    docx_document.add_heading(document.title or "문서", level=0)
    for block in document.blocks:
        if block.kind == "heading":
            docx_document.add_heading(block.markdown, level=1)
        elif block.kind == "paragraph":
            for line in block.markdown.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                bullet = re.match(r"^[-*]\s+(.+)$", stripped)
                if bullet:
                    docx_document.add_paragraph(bullet.group(1), style="List Bullet")
                else:
                    docx_document.add_paragraph(stripped)
        elif block.kind == "table":
            _docx_add_table(docx_document, _rows_as_text(block.rows))
        elif block.kind == "chart" and block.chart:
            caption = docx_document.add_paragraph(
                f"그래프: {block.chart.title or '차트'} (내보내기에서는 표로 표시됩니다)"
            )
            caption.runs[0].italic = True
            _docx_add_table(docx_document, _chart_fallback_rows(block.chart))
    buffer = io.BytesIO()
    docx_document.save(buffer)
    return f"{_safe_filename(document.title)}.docx", buffer.getvalue()


def export_hwpx(document: GeneratedDocument) -> tuple[str, bytes, dict[str, Any]]:
    from services.drafting_service import export_markdown_to_hwpx_with_validation

    return export_markdown_to_hwpx_with_validation(render_markdown(document), document.title or "문서")


def export_pdf(document: GeneratedDocument) -> tuple[str, bytes, dict[str, Any]]:
    from services.pdf_export_service import convert_hwpx_bytes_to_pdf

    hwpx_filename, hwpx_content, hwpx_summary = export_hwpx(document)
    pdf_filename, pdf_content, pdf_summary = convert_hwpx_bytes_to_pdf(
        hwpx_content, document.title or "문서", source_filename=hwpx_filename
    )
    return pdf_filename, pdf_content, {"hwpx_validation": hwpx_summary, "pdf_validation": pdf_summary}
