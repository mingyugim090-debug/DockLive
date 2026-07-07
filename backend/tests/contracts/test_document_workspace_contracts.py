import os
import sys
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[3]
BACKEND = ROOT / "backend"
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(BACKEND))

os.environ.setdefault("MOCK_MODE", "true")

try:
    from core.errors import AnalysisError  # noqa: E402
    from models.schemas import InlineTransformRequest, ParsedTableCell  # noqa: E402
    from services import workspace_service  # noqa: E402
    from services.analyzer import build_analysis_result  # noqa: E402
    from services.block_transforms import (  # noqa: E402
        apply_transform,
        paragraph_to_table,
        rewrite_block,
        table_to_chart,
    )
    from services.blueprint_service import build_blueprint, generate_document  # noqa: E402
    from services.mock_data import get_mock_result  # noqa: E402
    from services.spreadsheet_ingestion import parse_spreadsheet  # noqa: E402
    from services.workspace_export import export_hwpx, render_html, render_markdown  # noqa: E402
    try:
        from services import excel_artifacts  # noqa: E402
    except (ImportError, ModuleNotFoundError):
        excel_artifacts = None
except ModuleNotFoundError as exc:  # pragma: no cover - local minimal Python fallback
    if exc.name not in {"pydantic", "httpx", "fitz"}:
        raise
    workspace_service = None

BUDGET_CSV = (
    "항목,1차년도,2차년도\n"
    "인건비,42000000,45000000\n"
    "장비비,\"18,000,000\",6000000\n"
).encode("utf-8")

TEXT_ONLY_CSV = "구분,내용\n대상,소상공인\n방법,온라인 접수\n".encode("utf-8")


def _cell(text, row, col):
    return ParsedTableCell(text=text, row_index=row, col_index=col)


def _demo_workspace():
    workspace = workspace_service.create_workspace("테스트 프로젝트")
    workspace = workspace_service.get_workspace(workspace.id)
    workspace.analysis = build_analysis_result(get_mock_result("business_plan"), "demo", "demo-notice")
    workspace.status = "analyzed"
    workspace_service.save_workspace(workspace)
    return workspace_service.add_file(workspace.id, BUDGET_CSV, "예산.csv")


class SpreadsheetIngestionContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")

    def test_csv_parsing_preserves_values_verbatim(self):
        data = parse_spreadsheet(BUDGET_CSV, "예산.csv")
        self.assertEqual(len(data.sheets), 1)
        sheet = data.sheets[0]
        self.assertEqual(sheet.headers, ["항목", "1차년도", "2차년도"])
        self.assertEqual(sheet.rows[0], ["인건비", "42000000", "45000000"])
        # Quoted comma numbers survive exactly as written in the file.
        self.assertEqual(sheet.rows[1][1], "18,000,000")
        self.assertEqual(data.warnings, [])

    def test_broken_xlsx_yields_warning_not_content(self):
        data = parse_spreadsheet(b"PK\x03\x04fakexlsx", "budget.xlsx")
        self.assertEqual(data.sheets, [])
        self.assertTrue(data.warnings)

    def test_xlsx_parsing_preserves_values_verbatim(self):
        try:
            import io

            import openpyxl
        except ImportError:
            self.skipTest("openpyxl is not installed in this Python environment")

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "예산"
        sheet.append(["항목", "1차년도", "2차년도"])
        sheet.append(["인건비", 42000000, 45000000])
        sheet.append(["장비비", 18000000, 6000000])
        buffer = io.BytesIO()
        workbook.save(buffer)

        data = parse_spreadsheet(buffer.getvalue(), "예산.xlsx")
        self.assertEqual(len(data.sheets), 1)
        parsed = data.sheets[0]
        self.assertEqual(parsed.headers, ["항목", "1차년도", "2차년도"])
        self.assertEqual(parsed.rows[0], ["인건비", "42000000", "45000000"])
        self.assertEqual(parsed.rows[1], ["장비비", "18000000", "6000000"])
        self.assertEqual(data.warnings, [])


class WorkspaceServiceContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")

    def test_unsupported_file_is_recorded_with_warning_not_error(self):
        workspace = workspace_service.create_workspace("t")
        workspace = workspace_service.add_file(workspace.id, b"binarydata", "발표자료.pptx")
        self.assertEqual(len(workspace.files), 1)
        self.assertEqual(workspace.files[0].file_kind, "unsupported")
        self.assertTrue(workspace.files[0].warnings)
        self.assertEqual(workspace.files[0].text, "")

    def test_csv_file_becomes_spreadsheet_kind(self):
        workspace = workspace_service.create_workspace("t")
        workspace = workspace_service.add_file(workspace.id, BUDGET_CSV, "예산.csv")
        self.assertEqual(workspace.files[0].file_kind, "spreadsheet")
        self.assertEqual(workspace.status, "files_added")


class BlueprintContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")

    def test_planned_visuals_reference_only_existing_assets(self):
        workspace = _demo_workspace()
        blueprint = build_blueprint(workspace)

        valid_refs = set()
        for project_file in workspace.files:
            if project_file.sheet_data:
                for index in range(len(project_file.sheet_data.sheets)):
                    valid_refs.add(f"{project_file.id}:sheet-{index}")
        planned = [plan for section in blueprint.sections for plan in section.planned_visuals]
        self.assertTrue(planned)
        for plan in planned:
            self.assertIn(plan.source_ref, valid_refs)

    def test_no_chart_plans_without_numeric_data(self):
        workspace = workspace_service.create_workspace("t")
        workspace = workspace_service.add_file(workspace.id, TEXT_ONLY_CSV, "안내.csv")
        blueprint = build_blueprint(workspace_service.get_workspace(workspace.id))
        chart_plans = [
            plan
            for section in blueprint.sections
            for plan in section.planned_visuals
            if plan.kind == "chart"
        ]
        self.assertEqual(chart_plans, [])
        self.assertTrue(blueprint.confirmation_required)


class GenerateContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")

    def test_table_blocks_copy_source_values_verbatim(self):
        workspace = _demo_workspace()
        build_blueprint(workspace)
        document = generate_document(workspace_service.get_workspace(workspace.id))

        table_blocks = [block for block in document.blocks if block.kind == "table"]
        self.assertTrue(table_blocks)
        texts = ["|".join(cell.text for cell in row) for row in table_blocks[0].rows]
        self.assertIn("항목|1차년도|2차년도", texts)
        self.assertIn("인건비|42000000|45000000", texts)

        chart_blocks = [block for block in document.blocks if block.kind == "chart"]
        self.assertTrue(chart_blocks)
        chart = chart_blocks[0].chart
        self.assertEqual(chart.labels[0], "인건비")
        self.assertEqual(chart.series[0].values[0], 42000000.0)
        # 18,000,000 parses with the comma stripped but the magnitude unchanged.
        self.assertEqual(chart.series[0].values[1], 18000000.0)

    def test_sections_without_evidence_get_placeholder_with_needs_input(self):
        workspace = workspace_service.create_workspace("t")
        workspace = workspace_service.add_file(workspace.id, TEXT_ONLY_CSV, "안내.csv")
        build_blueprint(workspace_service.get_workspace(workspace.id))
        document = generate_document(workspace_service.get_workspace(workspace.id))
        placeholders = [block for block in document.blocks if block.status == "needs_input"]
        self.assertTrue(placeholders)
        for block in placeholders:
            self.assertIn("입력해 주세요", block.markdown)


class TransformContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")

    def test_paragraph_to_table_parses_label_value_lines(self):
        rows = paragraph_to_table("사업 기간: 2026년 3월 ~ 11월\n- 예산: 500만원")
        texts = [[cell.text for cell in row] for row in rows]
        self.assertEqual(texts[0], ["항목", "내용"])
        self.assertIn(["사업 기간", "2026년 3월 ~ 11월"], texts)
        self.assertIn(["예산", "500만원"], texts)

    def test_paragraph_without_structure_raises(self):
        with self.assertRaises(AnalysisError):
            paragraph_to_table("구조 없는 자유 서술 문장입니다.")

    def test_table_to_chart_preserves_values(self):
        rows = [
            [_cell("항목", 0, 0), _cell("금액", 0, 1)],
            [_cell("인건비", 1, 0), _cell("42,000,000원", 1, 1)],
            [_cell("장비비", 2, 0), _cell("18000000", 2, 1)],
        ]
        chart = table_to_chart(rows)
        self.assertEqual(chart.labels, ["인건비", "장비비"])
        self.assertEqual(chart.series[0].values, [42000000.0, 18000000.0])

    def test_table_without_numeric_column_raises(self):
        rows = [
            [_cell("구분", 0, 0), _cell("내용", 0, 1)],
            [_cell("대상", 1, 0), _cell("소상공인", 1, 1)],
        ]
        with self.assertRaises(AnalysisError):
            table_to_chart(rows)

    def test_rewrite_mock_only_normalizes_whitespace(self):
        original = "  첫  줄 입니다.  \n\n  둘째   줄  "
        result = rewrite_block(original)
        self.assertEqual("".join(result.split()), "".join(original.split()))

    def test_apply_transform_roundtrip_persists(self):
        workspace = _demo_workspace()
        build_blueprint(workspace)
        document = generate_document(workspace_service.get_workspace(workspace.id))
        table_block = next(block for block in document.blocks if block.kind == "table")

        transformed = apply_transform(
            workspace.id, table_block.id, InlineTransformRequest(command="to_chart")
        )
        self.assertEqual(transformed.kind, "chart")
        self.assertIsNotNone(transformed.chart)
        # rows are kept for table-fallback exports.
        self.assertTrue(transformed.rows)

        reloaded = workspace_service.get_workspace(workspace.id)
        saved = next(block for block in reloaded.document.blocks if block.id == table_block.id)
        self.assertEqual(saved.kind, "chart")


class ExportContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")

    def test_markdown_renders_tables_and_chart_fallback(self):
        workspace = _demo_workspace()
        build_blueprint(workspace)
        generate_document(workspace_service.get_workspace(workspace.id))
        reloaded = workspace_service.get_workspace(workspace.id)

        markdown = render_markdown(reloaded.document)
        self.assertIn("| 항목 | 1차년도 | 2차년도 |", markdown)
        self.assertIn("그래프:", markdown)
        self.assertIn("내보내기에서는 표로 표시됩니다", markdown)

        html = render_html(reloaded.document)
        self.assertIn("<table>", html)
        self.assertIn("인건비", html)

    def test_docx_export_preserves_title_tables_and_chart_fallback(self):
        try:
            import io as io_module

            from docx import Document as DocxDocument
        except ImportError:
            self.skipTest("python-docx is not installed in this Python environment")

        from services.workspace_export import export_docx

        workspace = _demo_workspace()
        build_blueprint(workspace)
        generate_document(workspace_service.get_workspace(workspace.id))
        reloaded = workspace_service.get_workspace(workspace.id)

        filename, content = export_docx(reloaded.document)
        self.assertTrue(filename.endswith(".docx"))
        self.assertTrue(content.startswith(b"PK"))

        parsed = DocxDocument(io_module.BytesIO(content))
        all_text = "\n".join(p.text for p in parsed.paragraphs)
        self.assertIn(reloaded.document.title, all_text)
        table_texts = [cell.text for table in parsed.tables for row in table.rows for cell in row.cells]
        self.assertIn("인건비", table_texts)
        self.assertIn("42000000", table_texts)
        # Chart blocks appear as fallback tables, so at least two tables exist.
        self.assertGreaterEqual(len(parsed.tables), 2)

    def test_hwpx_export_summary_marks_chart_fallback(self):
        from services import drafting_service

        original_export = drafting_service.export_markdown_to_hwpx_with_validation
        drafting_service.export_markdown_to_hwpx_with_validation = lambda markdown, title: (
            f"{title}.hwpx",
            b"PK\x03\x04hwpx",
            {"warnings": [], "validation_passed": True},
        )
        try:
            workspace = _demo_workspace()
            build_blueprint(workspace)
            generate_document(workspace_service.get_workspace(workspace.id))
            reloaded = workspace_service.get_workspace(workspace.id)

            chart_blocks = [block for block in reloaded.document.blocks if block.kind == "chart"]
            self.assertTrue(chart_blocks)

            filename, content, summary = export_hwpx(reloaded.document)
        finally:
            drafting_service.export_markdown_to_hwpx_with_validation = original_export

        self.assertTrue(filename.endswith(".hwpx"))
        self.assertTrue(content.startswith(b"PK"))
        warnings = " ".join(summary.get("warnings", []))
        self.assertIn("chart", warnings.lower())
        self.assertIn("source table", warnings.lower())
        self.assertTrue(summary.get("chart_fallback_used"))
        self.assertGreaterEqual(summary.get("chart_fallback_count", 0), 1)


class ExcelArtifactContractTests(unittest.TestCase):
    def setUp(self):
        if workspace_service is None:
            self.skipTest("backend dependencies are not installed in this Python environment")
        if excel_artifacts is None:
            self.fail("services.excel_artifacts module is required for Excel artifact generation")

    def test_workbook_plan_uses_source_refs_and_marks_missing_fields(self):
        workspace = _demo_workspace()
        plan = excel_artifacts.build_workbook_plan(workspace)

        self.assertEqual(plan.artifact_kind, "excel")
        self.assertEqual(
            [sheet.name for sheet in plan.sheets],
            ["\ub300\uc2dc\ubcf4\ub4dc", "\uc81c\ucd9c\uc11c\ub958", "\ucc28\ud2b8", "\uc6d0\ubb38\uadfc\uac70"],
        )
        table_refs = [table.source_ref for sheet in plan.sheets for table in sheet.tables]
        self.assertTrue(table_refs)
        self.assertTrue(all(ref for ref in table_refs))
        self.assertTrue(plan.confirmation_required)
        self.assertTrue(any("source" in item.lower() or "\uadfc\uac70" in item for item in plan.confirmation_required))

    def test_generate_excel_artifact_creates_xlsx_dashboard_with_chart(self):
        try:
            import openpyxl
        except ImportError:
            self.skipTest("openpyxl is not installed in this Python environment")

        workspace = _demo_workspace()
        artifact = excel_artifacts.generate_excel_artifact(workspace)
        reloaded = workspace_service.get_workspace(workspace.id)

        self.assertEqual(artifact.kind, "excel")
        self.assertTrue(artifact.filename.endswith(".xlsx"))
        self.assertTrue(artifact.storage_path)
        self.assertEqual(reloaded.artifacts[0].id, artifact.id)

        workbook = openpyxl.load_workbook(artifact.storage_path)
        try:
            self.assertEqual(
                workbook.sheetnames,
                ["\ub300\uc2dc\ubcf4\ub4dc", "\uc81c\ucd9c\uc11c\ub958", "\ucc28\ud2b8", "\uc6d0\ubb38\uadfc\uac70"],
            )
            dashboard = workbook["\ub300\uc2dc\ubcf4\ub4dc"]
            self.assertEqual(dashboard["A1"].value, workspace.analysis.title)
            self.assertTrue(workbook["\ucc28\ud2b8"]._charts)
        finally:
            workbook.close()

    def test_sync_excel_artifact_records_user_edit_snapshot(self):
        workspace = _demo_workspace()
        artifact = excel_artifacts.generate_excel_artifact(workspace)

        synced = excel_artifacts.sync_excel_artifact(workspace.id, artifact.id)

        self.assertEqual(synced.sync_state.status, "synced")
        self.assertTrue(synced.sync_state.snapshot)
        self.assertEqual(synced.sync_state.snapshot["source"], "user_edit")

    def test_desktop_helper_open_result_updates_artifact_state(self):
        workspace = _demo_workspace()
        artifact = excel_artifacts.generate_excel_artifact(workspace)

        calls = []

        def fake_helper(command, path, previous_mtime=0.0):
            calls.append((command, path, previous_mtime))
            return {
                "status": "opened",
                "path": path,
                "last_opened_at": "2026-07-07T01:00:00Z",
                "last_mtime": 123.5,
            }

        opened = excel_artifacts.open_excel_artifact(
            workspace.id,
            artifact.id,
            helper_runner=fake_helper,
        )

        self.assertEqual(calls[0][0], "open")
        self.assertEqual(calls[0][1], artifact.storage_path)
        self.assertEqual(opened.sync_state.status, "opened")
        self.assertEqual(opened.sync_state.last_opened_at, "2026-07-07T01:00:00Z")

    def test_desktop_helper_sync_result_records_user_edit_snapshot(self):
        workspace = _demo_workspace()
        artifact = excel_artifacts.generate_excel_artifact(workspace)
        artifact.sync_state.last_mtime = 100.0
        workspace.artifacts = [artifact]
        workspace_service.save_workspace(workspace)

        calls = []

        def fake_helper(command, path, previous_mtime=0.0):
            calls.append((command, path, previous_mtime))
            return {
                "status": "synced",
                "path": path,
                "last_synced_at": "2026-07-07T01:05:00Z",
                "last_mtime": 150.0,
                "snapshot": {"source": "user_edit", "sheets": {"dashboard": [["edited"]]}},
                "warnings": ["manual save detected"],
            }

        synced = excel_artifacts.sync_excel_artifact(
            workspace.id,
            artifact.id,
            helper_runner=fake_helper,
        )

        self.assertEqual(calls[0], ("watch-once", artifact.storage_path, 100.0))
        self.assertEqual(synced.sync_state.status, "synced")
        self.assertEqual(synced.sync_state.last_synced_at, "2026-07-07T01:05:00Z")
        self.assertEqual(synced.sync_state.last_mtime, 150.0)
        self.assertEqual(synced.sync_state.snapshot["sheets"]["dashboard"][0][0], "edited")
        self.assertEqual(synced.sync_state.warnings, ["manual save detected"])


if __name__ == "__main__":
    unittest.main()
