"""소스 문서 파싱 도구 (Phase 4에서 본격 구현). HWPX 규약은 hwpx-pipeline 스킬 참조."""
from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


def _ok(data) -> dict:
    return {"ok": True, "data": data}


def _err(msg: str) -> dict:
    return {"ok": False, "error": msg}


def list_files(dir_path: str) -> dict:
    p = Path(dir_path)
    if not p.is_dir():
        return _err(f"폴더가 없음: {dir_path}")
    files = sorted(f.name for f in p.iterdir() if f.is_file())
    return _ok(files)


MAX_PARAGRAPHS = 500


def read_document(path: str) -> dict:
    p = Path(path)
    if not p.exists():
        return _err(f"파일이 없음: {path}")
    suffix = p.suffix.lower()
    if suffix == ".hwpx":
        return _read_hwpx(p)
    if suffix == ".docx":
        return _read_docx(p)
    if suffix == ".pdf":
        return _read_pdf(p)
    return _err(f"지원하지 않는 형식: {suffix}")


def _read_docx(p: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        return _err("python-docx 미설치. pip install python-docx")
    try:
        document = Document(str(p))
        paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append(" | ".join(cells))
        return _ok({"file": p.name, "paragraphs": paragraphs[:MAX_PARAGRAPHS]})
    except Exception as ex:
        return _err(f"DOCX 파싱 실패: {ex}")


def _read_pdf(p: Path) -> dict:
    try:
        from pypdf import PdfReader
    except ImportError:
        return _err("pypdf 미설치. pip install pypdf")
    try:
        reader = PdfReader(str(p))
        paragraphs: list[str] = []
        for page in reader.pages:
            for line in (page.extract_text() or "").splitlines():
                if line.strip():
                    paragraphs.append(line.strip())
        if not paragraphs:
            return _err("PDF에서 텍스트를 추출하지 못함 (스캔본/이미지 PDF일 수 있음)")
        return _ok({"file": p.name, "paragraphs": paragraphs[:MAX_PARAGRAPHS]})
    except Exception as ex:
        return _err(f"PDF 파싱 실패: {ex}")


def relevant_excerpt(paragraphs: list[str], request: str, max_chars: int = 6000) -> str:
    """긴 문서에서 요청과 관련된 문단만 골라 컨텍스트 예산 안에 담는다 (Phase 4).

    문단별로 요청 토큰(2자 이상)과의 겹침을 세고, 점수 높은 문단부터
    max_chars까지 담되 원문 순서를 유지한다. 요청이 비면 문서 앞부분을 담는다.
    """
    tokens = {t for t in request.split() if len(t) >= 2}

    def score(paragraph: str) -> int:
        return sum(1 for t in tokens if t in paragraph)

    if tokens and any(score(p) for p in paragraphs):
        ranked = sorted(range(len(paragraphs)), key=lambda i: score(paragraphs[i]), reverse=True)
        chosen: set[int] = set()
        budget = max_chars
        for index in ranked:
            cost = len(paragraphs[index]) + 1
            if score(paragraphs[index]) == 0 or cost > budget:
                continue
            chosen.add(index)
            budget -= cost
        selected = [paragraphs[i] for i in sorted(chosen)]
    else:
        selected, budget = [], max_chars
        for paragraph in paragraphs:
            cost = len(paragraph) + 1
            if cost > budget:
                break
            selected.append(paragraph)
            budget -= cost
    return "\n".join(selected)


def _read_hwpx(p: Path) -> dict:
    """HWPX(zip+xml)에서 본문 텍스트를 문단 단위로 추출."""
    try:
        paragraphs: list[str] = []
        with zipfile.ZipFile(p) as zf:
            sections = sorted(n for n in zf.namelist()
                              if n.startswith("Contents/section") and n.endswith(".xml"))
            for name in sections:
                root = ET.fromstring(zf.read(name))
                for para in root.iter():
                    if para.tag.endswith("}p"):
                        text = "".join(
                            t.text or "" for t in para.iter() if t.tag.endswith("}t")
                        )
                        if text.strip():
                            paragraphs.append(text.strip())
        return _ok({"file": p.name, "paragraphs": paragraphs[:500]})
    except Exception as ex:
        return _err(f"HWPX 파싱 실패: {ex}")
